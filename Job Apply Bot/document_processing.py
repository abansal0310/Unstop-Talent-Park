import docx
import logging
import re
import os
import pytest

def test_learn_from_document(bot):
    bot.learn_from_document('valid_document.docx')
    assert bot.user_profile != {}

# Set up logging
logging.basicConfig(filename='document_processing.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def learn_from_document(document_path):
    try:
        user_profile = {}
        doc = docx.Document(document_path)

        for paragraph in doc.paragraphs:
            if ":" in paragraph.text:
                key, value = paragraph.text.split(":", 1)
                user_profile[key.strip()] = value.strip()

        if not user_profile:
            logging.warning(f"No user profile information found in the document: {document_path}")

        logging.info(f"Successfully extracted user profile from the document: {document_path}")
        return user_profile
    except docx.opc.exceptions.PackageNotFoundError:
        logging.error(f"The provided document file '{document_path}' is not a valid Word document.")
    except FileNotFoundError:
        logging.error(f"The document file '{document_path}' was not found.")
    except Exception as e:
        logging.error(f"An error occurred while processing the document: {e}")
        return {}

def extract_skills_from_document(document_path):
    try:
        skills = []
        doc = docx.Document(document_path)

        for paragraph in doc.paragraphs:
            skills.extend(re.findall(r'\b\w+\b', paragraph.text))

        skills = list(set([skill.lower() for skill in skills if skill.isalpha()]))

        if not skills:
            logging.warning(f"No skills found in the document: {document_path}")

        logging.info(f"Successfully extracted skills from the document: {document_path}")
        return skills
    except docx.opc.exceptions.PackageNotFoundError:
        logging.error(f"The provided document file '{document_path}' is not a valid Word document.")
    except FileNotFoundError:
        logging.error(f"The document file '{document_path}' was not found.")
    except Exception as e:
        logging.error(f"An error occurred while processing the document: {e}")
        return []

def extract_education_from_document(document_path):
    try:
        education = []
        doc = docx.Document(document_path)

        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join([cell.text for cell in row.cells])
                if 'Education' in row_text:
                    for cell in row.cells:
                        education.extend([item.strip() for item in cell.text.strip().split('\n') if item.strip()])

        if not education:
            logging.warning(f"No education information found in the document: {document_path}")

        logging.info(f"Successfully extracted education information from the document: {document_path}")
        return education
    except docx.opc.exceptions.PackageNotFoundError:
        logging.error(f"The provided document file '{document_path}' is not a valid Word document.")
    except FileNotFoundError:
        logging.error(f"The document file '{document_path}' was not found.")
    except Exception as e:
        logging.error(f"An error occurred while processing the document: {e}")
        return []